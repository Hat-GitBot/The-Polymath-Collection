

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from scipy import stats
import warnings
warnings.filterwarnings('ignore')

st.set_page_config(page_title="Universal Data Analytics", page_icon="⚜️", layout="wide")

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

class UniversalAnalyzer:
    """Universal data analyzer that works with any CSV structure"""

    def __init__(self, df):
        self.df = df
        self.numerical_cols = []
        self.categorical_cols = []
        self.is_sufficient = self._check_sufficiency()
        self._classify_columns()

    def _check_sufficiency(self):
        """Check if data is sufficient for analysis"""
        if len(self.df) < 10:
            return False
        if len(self.df.columns) < 2:
            return False
        # Check if there's at least some variation
        varied_cols = sum(1 for col in self.df.columns if self.df[col].nunique() > 1)
        return varied_cols >= 2

    def _classify_columns(self):
        """Classify all columns as numerical or categorical"""
        for col in self.df.columns:
            try:
                # Skip columns that are mostly unique (likely IDs)
                uniqueness = self.df[col].nunique() / len(self.df)
                if uniqueness > 0.95:
                    continue

                # Check data type
                if pd.api.types.is_numeric_dtype(self.df[col]):
                    # Skip if it looks like an ID or index
                    if 'id' in col.lower() or uniqueness == 1.0:
                        continue
                    # Skip if it's a year/month/day/hour column
                    if col.lower() in ['year', 'month', 'day', 'hour', 'week', 'quarter']:
                        continue
                    # Skip boolean columns stored as 0/1
                    if self.df[col].nunique() == 2 and set(self.df[col].dropna().unique()).issubset({0, 1, True, False}):
                        continue
                    if self.df[col].nunique() > 1:  # Has variation
                        self.numerical_cols.append(col)
                elif pd.api.types.is_bool_dtype(self.df[col]):
                    # Skip pure boolean columns
                    continue
                else:
                    # Categorical - but not if too many unique values
                    if uniqueness < 0.8:
                        self.categorical_cols.append(col)
            except Exception as e:
                # Skip problematic columns
                continue

    def analyze(self):
        """Perform complete analysis"""
        if not self.is_sufficient:
            return {'sufficient': False, 'reason': 'Insufficient data'}

        results = {
            'sufficient': True,
            'basic_stats': self._basic_statistics(),
            'numerical_analysis': self._analyze_numerical(),
            'categorical_analysis': self._analyze_categorical(),
            'relationships': self._find_relationships(),
            'insights': self._generate_insights()
        }
        return results

    def _basic_statistics(self):
        """Get basic dataset statistics"""
        return {
            'total_rows': len(self.df),
            'total_columns': len(self.df.columns),
            'numerical_count': len(self.numerical_cols),
            'categorical_count': len(self.categorical_cols),
            'missing_values': self.df.isnull().sum().sum(),
            'memory_usage': self.df.memory_usage(deep=True).sum() / 1024**2
        }

    def _analyze_numerical(self):
        """Analyze all numerical columns"""
        analysis = {}
        for col in self.numerical_cols:
            data = self.df[col].dropna()
            if len(data) > 0:
                analysis[col] = {
                    'mean': data.mean(),
                    'median': data.median(),
                    'std': data.std(),
                    'min': data.min(),
                    'max': data.max(),
                    'q25': data.quantile(0.25),
                    'q75': data.quantile(0.75),
                    'range': data.max() - data.min(),
                    'cv': (data.std() / data.mean() * 100) if data.mean() != 0 else 0
                }
        return analysis

    def _analyze_categorical(self):
        """Analyze all categorical columns"""
        analysis = {}
        for col in self.categorical_cols:
            value_counts = self.df[col].value_counts()
            analysis[col] = {
                'unique_count': len(value_counts),
                'top_value': value_counts.index[0] if len(value_counts) > 0 else None,
                'top_count': value_counts.iloc[0] if len(value_counts) > 0 else 0,
                'top_percentage': (value_counts.iloc[0] / len(self.df) * 100) if len(value_counts) > 0 else 0,
                'distribution': value_counts.head(10).to_dict()
            }
        return analysis

    def _find_relationships(self):
        """Find correlations between numerical columns"""
        relationships = []
        if len(self.numerical_cols) >= 2:
            for i, col1 in enumerate(self.numerical_cols):
                for col2 in self.numerical_cols[i+1:]:
                    try:
                        # Get clean data for both columns
                        data1 = self.df[col1].dropna()
                        data2 = self.df[col2].dropna()

                        # Find common indices
                        common_idx = data1.index.intersection(data2.index)

                        if len(common_idx) > 10:  # Need at least 10 points
                            clean_data1 = self.df.loc[common_idx, col1]
                            clean_data2 = self.df.loc[common_idx, col2]

                            # Calculate correlation
                            corr = clean_data1.corr(clean_data2)

                            if not np.isnan(corr) and abs(corr) > 0.3:  # Only significant correlations
                                relationships.append({
                                    'col1': col1,
                                    'col2': col2,
                                    'correlation': float(corr),
                                    'strength': 'Strong' if abs(corr) > 0.7 else 'Moderate'
                                })
                    except Exception as e:
                        continue
        return relationships

    def _generate_insights(self):
        """Generate automatic insights"""
        insights = []

        # Numerical insights
        num_analysis = self._analyze_numerical()
        for col, stats in num_analysis.items():
            # High variability
            if stats['cv'] > 50:
                insights.append(f"**{col}** shows high variability (CV: {stats['cv']:.1f}%)")

            # Skewed data
            if stats['mean'] > stats['median'] * 1.2:
                insights.append(f"**{col}** is right-skewed (mean > median)")
            elif stats['mean'] < stats['median'] * 0.8:
                insights.append(f"**{col}** is left-skewed (mean < median)")

        # Categorical insights
        cat_analysis = self._analyze_categorical()
        for col, stats in cat_analysis.items():
            if stats['top_percentage'] > 70:
                insights.append(f"**{col}** is highly concentrated: '{stats['top_value']}' represents {stats['top_percentage']:.1f}%")
            elif stats['unique_count'] > 20:
                insights.append(f"**{col}** has high diversity with {stats['unique_count']} unique values")

        # Relationship insights
        relationships = self._find_relationships()
        for rel in relationships[:3]:  # Top 3
            insights.append(f"**{rel['strength']} {('positive' if rel['correlation'] > 0 else 'negative')} correlation** between {rel['col1']} and {rel['col2']} (r={rel['correlation']:.2f})")

        return insights

def create_universal_visualizations(analyzer, results):
    """Create visualizations for ANY dataset"""
    figs = []

    # 1. Dataset Overview Dashboard
    fig1, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 10))

    stats = results['basic_stats']

    # Overview metrics
    ax1.text(0.5, 0.7, f"{stats['total_rows']:,}", ha='center', va='center',
            fontsize=48, fontweight='bold', color='#3b82f6')
    ax1.text(0.5, 0.3, 'Total Records', ha='center', va='center', fontsize=16)
    ax1.axis('off')
    ax1.add_patch(plt.Rectangle((0.1, 0.1), 0.8, 0.8, fill=False, edgecolor='#3b82f6', linewidth=3))

    ax2.text(0.5, 0.7, f"{stats['total_columns']}", ha='center', va='center',
            fontsize=48, fontweight='bold', color='#10b981')
    ax2.text(0.5, 0.3, 'Total Columns', ha='center', va='center', fontsize=16)
    ax2.axis('off')
    ax2.add_patch(plt.Rectangle((0.1, 0.1), 0.8, 0.8, fill=False, edgecolor='#10b981', linewidth=3))

    # Column type breakdown
    categories = ['Numerical', 'Categorical']
    counts = [stats['numerical_count'], stats['categorical_count']]
    colors = ['#3b82f6', '#f59e0b']
    ax3.pie(counts, labels=categories, autopct='%1.1f%%', colors=colors, startangle=90)
    ax3.set_title('Column Types', fontsize=14, fontweight='bold')

    # Data quality
    quality = (1 - stats['missing_values'] / (stats['total_rows'] * stats['total_columns'])) * 100
    ax4.bar(['Data Quality'], [quality], color='#10b981' if quality > 90 else '#f59e0b', alpha=0.7, width=0.4)
    ax4.set_ylim(0, 100)
    ax4.set_ylabel('Percentage (%)', fontsize=12)
    ax4.text(0, quality + 2, f'{quality:.1f}%', ha='center', fontweight='bold', fontsize=16)
    ax4.set_title('Data Quality Score', fontsize=14, fontweight='bold')
    ax4.grid(axis='y', alpha=0.3)

    plt.suptitle('Dataset Overview', fontsize=18, fontweight='bold', y=0.98)
    plt.tight_layout()
    figs.append(('dataset_overview', fig1))

    # 2. Numerical Distributions (top 4 numerical columns)
    if analyzer.numerical_cols:
        n_plots = min(4, len(analyzer.numerical_cols))
        fig2, axes = plt.subplots(2, 2, figsize=(16, 10))
        axes = axes.flatten()

        for idx, col in enumerate(analyzer.numerical_cols[:n_plots]):
            data = analyzer.df[col].dropna()

            axes[idx].hist(data, bins=min(30, len(data)//5), color='steelblue',
                          alpha=0.7, edgecolor='black')
            axes[idx].axvline(data.mean(), color='red', linestyle='--', linewidth=2,
                            label=f'Mean: {data.mean():.2f}')
            axes[idx].axvline(data.median(), color='green', linestyle='--', linewidth=2,
                            label=f'Median: {data.median():.2f}')
            axes[idx].set_title(f'{col}', fontsize=12, fontweight='bold')
            axes[idx].set_xlabel('Value', fontsize=10)
            axes[idx].set_ylabel('Frequency', fontsize=10)
            axes[idx].legend(fontsize=9)
            axes[idx].grid(True, alpha=0.3)

        # Hide unused subplots
        for idx in range(n_plots, 4):
            axes[idx].axis('off')

        plt.suptitle('Numerical Distributions', fontsize=18, fontweight='bold')
        plt.tight_layout()
        figs.append(('numerical_distributions', fig2))

    # 3. Categorical Distributions (top 3 categorical columns)
    if analyzer.categorical_cols:
        n_cats = min(3, len(analyzer.categorical_cols))
        fig3, axes = plt.subplots(1, n_cats, figsize=(6*n_cats, 6))
        if n_cats == 1:
            axes = [axes]

        for idx, col in enumerate(analyzer.categorical_cols[:n_cats]):
            value_counts = analyzer.df[col].value_counts().head(10)
            colors = plt.cm.Set3(np.linspace(0, 1, len(value_counts)))

            bars = axes[idx].barh(range(len(value_counts)), value_counts.values,
                                 color=colors, alpha=0.8, edgecolor='black')
            axes[idx].set_yticks(range(len(value_counts)))
            axes[idx].set_yticklabels(value_counts.index, fontsize=9)
            axes[idx].invert_yaxis()
            axes[idx].set_xlabel('Count', fontsize=10)
            axes[idx].set_title(f'{col}', fontsize=12, fontweight='bold')
            axes[idx].grid(True, alpha=0.3, axis='x')

            # Add percentages
            total = len(analyzer.df)
            for i, (bar, val) in enumerate(zip(bars, value_counts.values)):
                pct = (val / total) * 100
                axes[idx].text(bar.get_width() + total*0.01, bar.get_y() + bar.get_height()/2,
                             f'{val} ({pct:.1f}%)', va='center', fontsize=9, fontweight='bold')

        plt.suptitle('Categorical Distributions', fontsize=18, fontweight='bold')
        plt.tight_layout()
        figs.append(('categorical_distributions', fig3))

    # 4. Correlation Heatmap (if multiple numerical columns)
    if len(analyzer.numerical_cols) >= 2:
        try:
            fig4, ax = plt.subplots(figsize=(12, 10))

            # Calculate correlation matrix with proper handling
            corr_data = analyzer.df[analyzer.numerical_cols].copy()
            corr_matrix = corr_data.corr()

            # Create mask for upper triangle using proper numpy operation
            mask = np.zeros_like(corr_matrix, dtype=bool)
            mask[np.triu_indices_from(mask, k=1)] = True

            sns.heatmap(corr_matrix, mask=mask, annot=True, fmt='.2f',
                       cmap='RdYlGn', center=0, square=True, linewidths=1,
                       cbar_kws={"shrink": 0.8}, ax=ax, vmin=-1, vmax=1)

            ax.set_title('Correlation Matrix', fontsize=16, fontweight='bold', pad=20)
            plt.tight_layout()
            figs.append(('correlation_heatmap', fig4))
        except Exception as e:
            # Skip heatmap if there's an error
            pass

    # 5. Box Plots for Numerical Columns (outlier detection)
    if analyzer.numerical_cols:
        n_plots = min(6, len(analyzer.numerical_cols))
        fig5, axes = plt.subplots(2, 3, figsize=(18, 10))
        axes = axes.flatten()

        for idx, col in enumerate(analyzer.numerical_cols[:n_plots]):
            data = analyzer.df[col].dropna()

            bp = axes[idx].boxplot([data], vert=True, patch_artist=True, widths=0.5)
            bp['boxes'][0].set_facecolor('lightblue')
            bp['boxes'][0].set_edgecolor('darkblue')
            bp['medians'][0].set_color('red')
            bp['medians'][0].set_linewidth(2)

            axes[idx].set_title(f'{col}', fontsize=11, fontweight='bold')
            axes[idx].set_ylabel('Value', fontsize=10)
            axes[idx].grid(True, alpha=0.3, axis='y')

            # Add statistics text
            stats_text = f'Mean: {data.mean():.2f}\nMedian: {data.median():.2f}\nStd: {data.std():.2f}'
            axes[idx].text(1.15, data.median(), stats_text, fontsize=9,
                         bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))

        # Hide unused subplots
        for idx in range(n_plots, 6):
            axes[idx].axis('off')

        plt.suptitle('Box Plots (Outlier Detection)', fontsize=18, fontweight='bold')
        plt.tight_layout()
        figs.append(('box_plots', fig5))

    # 6. Scatter Plots (top correlations)
    relationships = results.get('relationships', [])
    if relationships:
        try:
            n_plots = min(4, len(relationships))
            fig6, axes = plt.subplots(2, 2, figsize=(16, 12))
            axes = axes.flatten()

            for idx, rel in enumerate(relationships[:n_plots]):
                # Get data for both columns
                x_data = analyzer.df[rel['col1']].dropna()
                y_data = analyzer.df[rel['col2']].dropna()

                # Align data on common indices
                common_idx = x_data.index.intersection(y_data.index)
                x_data = analyzer.df.loc[common_idx, rel['col1']]
                y_data = analyzer.df.loc[common_idx, rel['col2']]

                if len(x_data) > 0:
                    axes[idx].scatter(x_data, y_data, alpha=0.6, s=50, color='steelblue', edgecolor='black')

                    # Add trend line
                    try:
                        z = np.polyfit(x_data, y_data, 1)
                        p = np.poly1d(z)
                        x_sorted = np.sort(x_data)
                        axes[idx].plot(x_sorted, p(x_sorted), "r--", linewidth=2, label='Trend')
                    except:
                        pass

                    axes[idx].set_xlabel(rel['col1'], fontsize=10)
                    axes[idx].set_ylabel(rel['col2'], fontsize=10)
                    axes[idx].set_title(f"{rel['col1']} vs {rel['col2']}\n{rel['strength']} Correlation (r={rel['correlation']:.2f})",
                                      fontsize=11, fontweight='bold')
                    axes[idx].legend()
                    axes[idx].grid(True, alpha=0.3)

            # Hide unused subplots
            for idx in range(n_plots, 4):
                axes[idx].axis('off')

            plt.suptitle('Relationship Analysis', fontsize=18, fontweight='bold')
            plt.tight_layout()
            figs.append(('scatter_plots', fig6))
        except Exception as e:
            # Skip scatter plots if there's an error
            pass

    return figs

def generate_universal_narrative(analyzer, results):
    """Generate narrative for any dataset"""
    stats = results['basic_stats']
    num_analysis = results['numerical_analysis']
    cat_analysis = results['categorical_analysis']
    insights = results['insights']

    narrative = f"""## Executive Summary

**Dataset Overview:** This analysis examines a dataset containing **{stats['total_rows']:,} records** across **{stats['total_columns']} dimensions**, comprising **{stats['numerical_count']} numerical variables** and **{stats['categorical_count']} categorical variables**.

**Data Quality:** {(1 - stats['missing_values'] / (stats['total_rows'] * stats['total_columns']) * 100):.1f}% complete with {stats['missing_values']:,} missing values.

---

### Key Numerical Findings

"""

    # Numerical analysis
    if num_analysis:
        for col, stats_dict in list(num_analysis.items())[:5]:
            narrative += f"""
**{col}:**
- Range: {stats_dict['min']:.2f} to {stats_dict['max']:.2f}
- Average: {stats_dict['mean']:.2f} (Median: {stats_dict['median']:.2f})
- Variability: ±{stats_dict['std']:.2f} (CV: {stats_dict['cv']:.1f}%)
"""

    # Categorical analysis
    if cat_analysis:
        narrative += "\n### Key Categorical Findings\n"
        for col, stats_dict in list(cat_analysis.items())[:5]:
            narrative += f"""
**{col}:**
- {stats_dict['unique_count']} unique categories
- Most common: '{stats_dict['top_value']}' ({stats_dict['top_percentage']:.1f}% of records)
"""

    # Insights
    if insights:
        narrative += "\n### Automated Insights\n"
        for insight in insights[:8]:
            narrative += f"- {insight}\n"

    # Relationships
    relationships = results.get('relationships', [])
    if relationships:
        narrative += "\n### Identified Relationships\n"
        for rel in relationships[:5]:
            direction = "positive" if rel['correlation'] > 0 else "negative"
            narrative += f"- **{rel['strength']} {direction} relationship** between {rel['col1']} and {rel['col2']} (correlation: {rel['correlation']:.2f})\n"

    return narrative

def generate_universal_recommendations(results):
    """Generate recommendations for any dataset"""
    recommendations = []
    stats = results['basic_stats']
    num_analysis = results['numerical_analysis']
    cat_analysis = results['categorical_analysis']
    relationships = results.get('relationships', [])

    # Data quality recommendations
    if stats['missing_values'] > stats['total_rows'] * 0.05:
        recommendations.append({
            'priority': '🟡 HIGH',
            'title': 'Address Missing Data',
            'action': f'Investigate and resolve {stats["missing_values"]:,} missing values ({stats["missing_values"]/(stats["total_rows"]*stats["total_columns"])*100:.1f}% of dataset)',
            'impact': 'Improve data quality and analysis reliability'
        })

    # High variability
    high_var_cols = [col for col, s in num_analysis.items() if s['cv'] > 50]
    if high_var_cols:
        recommendations.append({
            'priority': '🟢 MEDIUM',
            'title': 'Investigate High Variability',
            'action': f'Examine high variability in: {", ".join(high_var_cols[:3])}',
            'impact': 'Understand causes of inconsistency'
        })

    # Concentration issues
    concentrated_cats = [col for col, s in cat_analysis.items() if s['top_percentage'] > 80]
    if concentrated_cats:
        recommendations.append({
            'priority': '🟢 MEDIUM',
            'title': 'Review Data Distribution',
            'action': f'High concentration detected in: {", ".join(concentrated_cats[:3])}',
            'impact': 'Consider if this represents actual pattern or data collection bias'
        })

    # Strong relationships
    if relationships:
        strong_rels = [r for r in relationships if abs(r['correlation']) > 0.7]
        if strong_rels:
            recommendations.append({
                'priority': '🟢 MEDIUM',
                'title': 'Leverage Strong Correlations',
                'action': f'Explore causal relationships: {strong_rels[0]["col1"]} ↔ {strong_rels[0]["col2"]}',
                'impact': 'Identify predictive factors and dependencies'
            })

    # Sample size
    if stats['total_rows'] < 100:
        recommendations.append({
            'priority': '🔴 CRITICAL',
            'title': 'Increase Sample Size',
            'action': f'Collect more data - current {stats["total_rows"]} records may limit statistical reliability',
            'impact': 'Enable more robust analysis and confident conclusions'
        })

    if not recommendations:
        recommendations.append({
            'priority': '✅ INFO',
            'title': 'Data Quality Good',
            'action': 'Continue monitoring data patterns and trends',
            'impact': 'Maintain current data collection standards'
        })

    return recommendations

def generate_word_report(analyzer, results, narrative, recommendations, visualizations):
    """Generate Word report"""
    buffer = io.BytesIO()
    doc = Document()

    # Title
    title = doc.add_heading('Data Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(59, 130, 246)

    date_para = doc.add_paragraph(f'Generated: {pd.Timestamp.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Narrative
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(narrative)

    doc.add_page_break()

    # Statistical Summary
    doc.add_heading('Statistical Summary', level=1)

    num_analysis = results['numerical_analysis']
    if num_analysis:
        doc.add_heading('Numerical Variables', level=2)

        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'

        headers = ['Variable', 'Mean', 'Median', 'Std Dev', 'Min', 'Max', 'CV%']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        for col, stats in num_analysis.items():
            row = table.add_row()
            row.cells[0].text = col
            row.cells[1].text = f"{stats['mean']:.2f}"
            row.cells[2].text = f"{stats['median']:.2f}"
            row.cells[3].text = f"{stats['std']:.2f}"
            row.cells[4].text = f"{stats['min']:.2f}"
            row.cells[5].text = f"{stats['max']:.2f}"
            row.cells[6].text = f"{stats['cv']:.1f}"

    cat_analysis = results['categorical_analysis']
    if cat_analysis:
        doc.add_heading('Categorical Variables', level=2)

        for col, stats in cat_analysis.items():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{col}: ").bold = True
            p.add_run(f"{stats['unique_count']} categories, most common is '{stats['top_value']}' ({stats['top_percentage']:.1f}%)")

    doc.add_page_break()

    # Recommendations
    doc.add_heading('Recommendations', level=1)

    for rec in recommendations:
        doc.add_heading(f"{rec['priority']}: {rec['title']}", level=2)

        rec_table = doc.add_table(rows=2, cols=2)
        rec_table.style = 'Light List Accent 1'

        rec_table.rows[0].cells[0].text = 'Action'
        rec_table.rows[0].cells[1].text = rec['action']
        rec_table.rows[1].cells[0].text = 'Expected Impact'
        rec_table.rows[1].cells[1].text = rec['impact']

        doc.add_paragraph()

    doc.add_page_break()

    # Visualizations
    doc.add_heading('Visualizations', level=1)

    for title, fig in visualizations:
        doc.add_heading(title.replace('_', ' ').title(), level=2)

        img_buffer = io.BytesIO()
        fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6.5))
        doc.add_paragraph()
        plt.close(fig)

    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.markdown('<h1 class="main-header"><span style="-webkit-text-fill-color: initial">⚜️</span> Universal Data Analytics Platform</h1>', unsafe_allow_html=True)
    st.markdown("### Automatically analyze ANY CSV file - No configuration needed!")

    with st.sidebar:
        st.markdown("## 🎯 Features")
        st.markdown("""
        **✅ Works with ANY CSV**
        - Automatic column detection
        - Smart data classification
        - Universal visualizations

        **📊 Comprehensive Analysis**
        - Numerical distributions
        - Categorical breakdowns
        - Correlation analysis
        - Outlier detection

        **📈 Multiple Visualizations**
        - Overview dashboard
        - Distribution plots
        - Box plots
        - Heatmaps
        - Scatter plots

        **📄 Professional Reports**
        - Executive summaries
        - Statistical tables
        - Actionable recommendations
        - Publication-quality charts
        """)

    uploaded_file = st.file_uploader("📂 Upload Your CSV File", type="csv")

    if uploaded_file is not None:
        try:
            df = pd.read_csv(uploaded_file)

            st.markdown("---")
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric("📊 Records", f"{len(df):,}")
            with col2:
                st.metric("📋 Columns", len(df.columns))
            with col3:
                quality = (1 - df.isnull().sum().sum()/(len(df)*len(df.columns)))*100
                st.metric("✅ Quality", f"{quality:.1f}%")
            with col4:
                memory = df.memory_usage(deep=True).sum() / 1024**2
                st.metric("💾 Size", f"{memory:.1f} MB")

            with st.expander("👀 Preview Data"):
                st.dataframe(df.head(20), use_container_width=True)

            st.markdown("---")

            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("🚀 Analyze Data", type="primary", use_container_width=True):

                    with st.spinner("🔍 Analyzing your data..."):
                        progress = st.progress(0)
                        status = st.empty()

                        status.text("Initializing analyzer...")
                        progress.progress(10)
                        analyzer = UniversalAnalyzer(df)

                        # Check sufficiency
                        if not analyzer.is_sufficient:
                            progress.progress(100)
                            status.text("Analysis complete")
                            st.error("❌ **Insufficient Data**")
                            st.warning("""
                            This dataset does not have enough data for meaningful analysis.

                            **Requirements:**
                            - At least 10 records
                            - At least 2 columns
                            - At least 2 columns with variation

                            **Current Dataset:**
                            - Records: {len(df)}
                            - Columns: {len(df.columns)}

                            Please provide a larger dataset for comprehensive analysis.
                            """)
                        else:
                            # Perform analysis
                            status.text("Analyzing data patterns...")
                            progress.progress(30)

                            try:
                                results = analyzer.analyze()
                            except Exception as e:
                                st.error(f"Error during analysis: {str(e)}")
                                import traceback
                                with st.expander("See full error"):
                                    st.code(traceback.format_exc())
                                return

                            status.text("Generating insights...")
                            progress.progress(50)

                            try:
                                narrative = generate_universal_narrative(analyzer, results)
                            except Exception as e:
                                st.error(f"Error generating narrative: {str(e)}")
                                narrative = "Error generating narrative"

                            status.text("Creating recommendations...")
                            progress.progress(65)

                            try:
                                recommendations = generate_universal_recommendations(results)
                            except Exception as e:
                                st.error(f"Error generating recommendations: {str(e)}")
                                recommendations = []

                            status.text("Creating visualizations...")
                            progress.progress(80)

                            try:
                                visualizations = create_universal_visualizations(analyzer, results)
                            except Exception as e:
                                st.error(f"❌ Error creating visualizations: {str(e)}")
                                import traceback
                                with st.expander("See visualization error details"):
                                    st.code(traceback.format_exc())
                                    st.write("**Detected columns:**")
                                    st.write(f"Numerical: {analyzer.numerical_cols}")
                                    st.write(f"Categorical: {analyzer.categorical_cols}")
                                visualizations = []

                            status.text("Generating report...")
                            progress.progress(95)

                            try:
                                report = generate_word_report(analyzer, results, narrative,
                                                             recommendations, visualizations)
                            except Exception as e:
                                st.error(f"Error generating report: {str(e)}")
                                report = None

                            progress.progress(100)
                            status.text("✅ Analysis complete!")

                            st.success("🎉 Analysis Complete!")

                            # Download button
                            if report:
                                st.download_button(
                                    label="📄 Download Complete Report (Word)",
                                    data=report.getvalue(),
                                    file_name=f"data_analysis_report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )

                            st.markdown("---")

                            # Display narrative
                            st.markdown(narrative)

                            st.markdown("---")

                            # Show key statistics
                            st.subheader("📊 Key Statistics")

                            stats = results['basic_stats']
                            col1, col2, col3, col4 = st.columns(4)

                            with col1:
                                st.metric("Numerical Variables", stats['numerical_count'])
                            with col2:
                                st.metric("Categorical Variables", stats['categorical_count'])
                            with col3:
                                quality = (1 - stats['missing_values']/(stats['total_rows']*stats['total_columns']))*100
                                st.metric("Data Quality", f"{quality:.1f}%")
                            with col4:
                                st.metric("Missing Values", f"{stats['missing_values']:,}")

                            # Show insights
                            if results['insights']:
                                st.markdown("---")
                                st.subheader("💡 Key Insights")

                                for insight in results['insights'][:10]:
                                    st.info(insight)

                            # Show recommendations
                            st.markdown("---")
                            st.subheader("🎯 Recommendations")

                            for rec in recommendations:
                                with st.expander(f"{rec['priority']}: {rec['title']}", expanded=True):
                                    st.markdown(f"**Action:** {rec['action']}")
                                    st.markdown(f"**Impact:** {rec['impact']}")

                            # Show visualizations
                            st.markdown("---")
                            st.subheader("📈 Visualizations")

                            for title, fig in visualizations:
                                st.pyplot(fig)
                                plt.close(fig)

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.info("Please ensure your CSV is properly formatted")
            import traceback
            with st.expander("See error details"):
                st.code(traceback.format_exc())

    else:
        # Landing
        st.markdown("---")
        st.info("👆 Upload your CSV file to begin automatic analysis")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("""
            ### 🔄 Universal Analysis
            - Works with ANY CSV structure
            - No configuration needed
            - Automatic column detection
            - Smart data classification
            """)

        with col2:
            st.markdown("""
            ### 📊 Rich Visualizations
            - Distribution plots
            - Correlation heatmaps
            - Box plots for outliers
            - Scatter plots
            - Category breakdowns
            """)

        with col3:
            st.markdown("""
            ### 📝 Professional Reports
            - Executive summaries
            - Statistical analysis
            - Automated insights
            - Actionable recommendations
            - Downloadable Word docs
            """)

        st.markdown("---")

        st.markdown("""
        ### 📋 Supported Data Types

        This platform automatically analyzes:

        - **Customer Data**: satisfaction scores, demographics, behavior
        - **Sales Data**: revenue, transactions, conversion rates
        - **Operational Data**: tickets, resolution times, costs
        - **Marketing Data**: campaigns, engagement, ROI
        - **Financial Data**: budgets, expenses, profits
        - **Survey Data**: responses, ratings, feedback
        - **A/B Test Data**: variants, metrics, results
        - **Any structured CSV data!**

        **No special formatting required** - just upload and analyze!
        """)

if __name__ == "__main__":
    main()
